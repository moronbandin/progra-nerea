#!/usr/bin/env python3
"""Extract structural facts from the three source DOCX files."""

from __future__ import annotations

import json
import re
import sys
from collections import Counter
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


def paragraph_record(paragraph):
    text = " ".join(paragraph.text.split())
    props = paragraph._p.pPr
    page_break = bool(props is not None and props.find(qn("w:pageBreakBefore")) is not None)
    return {
        "text": text,
        "style": paragraph.style.name if paragraph.style else None,
        "pageBreakBefore": page_break,
        "boldRuns": sum(1 for run in paragraph.runs if run.bold),
        "italicRuns": sum(1 for run in paragraph.runs if run.italic),
    }


def analyze(path: Path) -> dict:
    document = Document(path)
    paragraphs = [paragraph_record(p) for p in document.paragraphs if p.text.strip()]
    tables = []
    for index, table in enumerate(document.tables, 1):
        rows = []
        for row in table.rows:
            rows.append([" ".join(cell.text.split()) for cell in row.cells])
        tables.append(
            {
                "index": index,
                "rows": len(table.rows),
                "columns": max((len(row.cells) for row in table.rows), default=0),
                "preview": rows[:3],
            }
        )

    section_pattern = re.compile(r"^(?P<number>1[0-2]|[1-9])\.\s+(?P<title>.+)$")
    session_pattern = re.compile(r"^Sesi[oó]n\s+(?P<number>\d+)\s*[:—-]\s*(?P<title>.+)", re.I)
    activity_pattern = re.compile(r"^Actividad\s+(?P<number>\d+)\s*[:—-]\s*(?P<title>.+)", re.I)

    sections = []
    sessions = []
    activities = []
    for position, paragraph in enumerate(paragraphs):
        text = paragraph["text"]
        if match := section_pattern.match(text):
            sections.append({"position": position, **match.groupdict()})
        if match := session_pattern.match(text):
            sessions.append({"position": position, **match.groupdict()})
        if match := activity_pattern.match(text):
            activities.append({"position": position, **match.groupdict()})

    rels = document.part.rels.values()
    images = [
        {
            "target": str(rel.target_ref),
            "contentType": getattr(rel.target_part, "content_type", None),
        }
        for rel in rels
        if "image" in rel.reltype
    ]

    styles = Counter(p["style"] for p in paragraphs)
    return {
        "file": path.name,
        "paragraphCount": len(paragraphs),
        "tableCount": len(tables),
        "imageCount": len(images),
        "sectionCount": len(document.sections),
        "styles": dict(styles.most_common()),
        "sections": sections,
        "sessions": sessions,
        "activities": activities,
        "tables": tables,
        "images": images,
        "paragraphs": paragraphs,
    }


def main() -> None:
    paths = [Path(arg) for arg in sys.argv[1:]]
    print(json.dumps([analyze(path) for path in paths], ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
